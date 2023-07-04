from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import csv
from datetime import datetime
from docx.shared import Pt
from dateutil.relativedelta import relativedelta
from docx2pdf import convert

# Open the CSV file
with open('[name].csv', 'r') as file:
    # Create a CSV reader object
    reader = csv.reader(file)

    # Skip the header row
    header = next(reader)

    # Read the remaining rows and store them in a list
    rows = [row for row in reader]

# Sort the rows by the title column
sorted_rows = sorted(rows, key=lambda row: row[6])  # Assuming the title column is at index 6

# Get the current date
current_date = datetime.now()

# Calculate the date for the last month
last_month_date = current_date - relativedelta(months=1)

# Get the last month and current year
last_month = last_month_date.strftime("%B")
current_year = current_date.year

# Create a new Word document
doc = Document()

# Add your information
header_text = "[Name]\n[Address Line 1]\n[Address Line 2]\n[Address Line 3]\n[City]\n[Postcode]"
header_paragraph = doc.add_paragraph()
header_paragraph.text = header_text
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add recipient information
recipient_text = "[Name]\n[Address Line 1]\n[Address Line 2]\n[City]\n[Postcode]"
recipient_paragraph = doc.add_paragraph()
recipient_paragraph.text = recipient_text

# Add title with the last month and current year
title_text = f"To: \n[Business] for {last_month} {current_year}"
title_paragraph = doc.add_paragraph()
title_paragraph.text = title_text

# Process the sorted data and add it to the document
previous_title = None
total_earnings = 0

for row in sorted_rows:
    date = row[1].split()[0]  # Extract only the date part
    duration = row[3]
    earnings = row[5]
    title = row[6].strip()  # Remove leading/trailing whitespace from the title

    # Check if the title has changed
    if title != previous_title:
        # Add a blank line between headers
        doc.add_paragraph()

        # Add the new title as a bold paragraph
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True

    # Add the data underneath the title
    data_paragraph = doc.add_paragraph()
    data_paragraph.add_run(date + " - " + duration + " hours - ").bold = False
    earnings_run = data_paragraph.add_run("£" + earnings)
    earnings_run.bold = True

    # Update the previous_title variable
    previous_title = title

    # Update the total earnings
    total_earnings += float(earnings)

# Add a paragraph for the total earnings
total_paragraph = doc.add_paragraph()
total_paragraph_run = total_paragraph.add_run("Total earnings: £" + str(total_earnings))
total_paragraph_run.bold = True
total_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Save the document and convert to pdf
doc.save("invoice.docx")
convert("invoice.docx")
